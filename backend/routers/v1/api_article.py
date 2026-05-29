import os
import uuid
import aiofiles
import traceback
import math
import hashlib
import zipfile
import tempfile
import shutil
import re
from datetime import datetime
from typing import Optional, List
from fastapi import APIRouter, Depends, Body, HTTPException, status, UploadFile, File, Query, Request
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, delete, and_, func
from sqlalchemy.orm import selectinload
from sqlalchemy.exc import IntegrityError

from dependencies import get_db, get_current_user, allow_admin_only, get_current_user_optional
from models.blog_models import Article, ArticleStatus, User, Category, Tag, UserRole, article_tag, ArticleLike
from schemas.article_schema import ArticleCreate, ArticleReviewAction, ArticleDetailOut

router = APIRouter()

IMAGE_STORAGE = "storage/images"
ARTICLE_STORAGE = "storage/articles"
_upload_counter: dict = {}


# --- 接口 1：POST /upload-image (图片上传) ---
@router.post("/upload-image", summary="图片上传接口")
async def upload_article_image(file: UploadFile = File(...), user: User = Depends(get_current_user)):
    now = datetime.now()
    user_id = user.id
    if user_id not in _upload_counter:
        _upload_counter[user_id] = []
    _upload_counter[user_id] = [t for t in _upload_counter[user_id] if (now - t).total_seconds() < 3600]
    if len(_upload_counter[user_id]) >= 15:
        raise HTTPException(status_code=429, detail="上传过于频繁，请1小时后再试")

    if not file.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="只能上传图片文件")
    ext = os.path.splitext(file.filename)[1]
    unique_name = f"{uuid.uuid4().hex}{ext}"
    os.makedirs(IMAGE_STORAGE, exist_ok=True)
    file_path = os.path.join(IMAGE_STORAGE, unique_name)
    try:
        content = await file.read()
        if len(content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="文件大小不能超过10MB")
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)
        _upload_counter[user_id].append(now)
        return {"url": f"/storage/images/{unique_name}"}
    except Exception:
        raise HTTPException(status_code=500, detail="文件保存失败")


# --- 接口 2：DELETE /upload-image (图片删除) ---
@router.delete("/upload-image", summary="图片删除接口")
async def delete_article_image(filename: str, user: User = Depends(get_current_user)):
    file_path = os.path.join(IMAGE_STORAGE, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="图片不存在")
    try:
        os.remove(file_path)
        return {"message": "图片已删除"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"删除文件时出错: {str(e)}")


# --- 接口 3：POST /autosave (加入乐观锁逻辑) ---
@router.post("/autosave", summary="自动保存/草稿更新")
async def autosave(article_in: ArticleCreate, user: User = Depends(get_current_user),
                   db: AsyncSession = Depends(get_db)):
    tags = []
    if article_in.tag_ids:
        tag_res = await db.execute(select(Tag).where(Tag.id.in_(article_in.tag_ids)))
        tags = tag_res.scalars().all()

    if article_in.id:
        res = await db.execute(select(Article).where(Article.id == article_in.id).options(selectinload(Article.tags)))
        article = res.scalars().first()

        if not article or (article.user_id != user.id and user.role != UserRole.ADMIN):
            raise HTTPException(status_code=403, detail="无权修改此文章")

        if article_in.version is not None and article_in.version != article.version:
            raise HTTPException(status_code=409, detail="文章已被他人修改，请刷新后重新编辑")

        article.title = article_in.title
        article.summary = article_in.summary
        if article_in.cover_image is not None:
            article.cover_image = article_in.cover_image
        article.category_id = article_in.category_id
        article.updated_at = datetime.now()
        article.tags = tags

        if article_in.content is not None:
            new_hash = hashlib.md5(article_in.content.encode()).hexdigest()
            if article.content_hash != new_hash:
                article.content = article_in.content
                article.content_hash = new_hash
        article.version += 1
    else:
        if (not article_in.title or not article_in.title.strip()) and not article_in.content:
            raise HTTPException(status_code=400, detail="标题和内容不能同时为空")

        content_hash = hashlib.md5(article_in.content.encode()).hexdigest() if article_in.content else ""
        article = Article(
            title=article_in.title, summary=article_in.summary, content=article_in.content,
            content_hash=content_hash, category_id=article_in.category_id, user_id=user.id,
            cover_image=article_in.cover_image,
            status=ArticleStatus.DRAFT, tags=tags, version=1
        )
        db.add(article)

    await db.commit()
    await db.refresh(article)
    return {"id": article.id, "version": article.version, "message": "已自动保存"}


# --- 接口 4：PUT /{article_id}/publish (发布文章) ---
@router.put("/{article_id}/publish", summary="正式发布文章")
async def publish_article(article_id: int, user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(Article).where(Article.id == article_id))
    article = res.scalars().first()

    if not article or (article.user_id != user.id and user.role != UserRole.ADMIN):
        raise HTTPException(status_code=403, detail="无权操作")

    if not article.title or not article.title.strip():
        raise HTTPException(status_code=400, detail="发布失败：标题不能为空")
    if not article.content or not article.content.strip():
        raise HTTPException(status_code=400, detail="发布失败：文章内容不能为空")

    if user.role == UserRole.ADMIN:
        article.status = ArticleStatus.PUBLISHED
        article.published_at = datetime.now()
    else:
        article.status = ArticleStatus.PENDING
        article.review_remark = None
        article.submitted_at = datetime.now()

    article.updated_at = datetime.now()
    await db.commit()
    return {"message": "发布成功" if user.role == UserRole.ADMIN else "已提交审核"}


# --- 接口 5：GET /{article_id} (获取详情) ---
@router.get("/{article_id}", response_model=ArticleDetailOut, summary="获取文章详情")
async def get_article_detail(article_id: int, user: Optional[User] = Depends(get_current_user_optional),
                             db: AsyncSession = Depends(get_db)):
    res = await db.execute(
        select(Article).where(Article.id == article_id).options(
            selectinload(Article.category), selectinload(Article.tags), selectinload(Article.author)
        )
    )
    article = res.scalars().first()
    if not article:
        raise HTTPException(status_code=404, detail="文章不存在")

    is_author = user and article.user_id == user.id
    is_admin = user and user.role == UserRole.ADMIN

    if article.deleted_at and not (is_author or is_admin):
        raise HTTPException(status_code=404, detail="文章已删除")

    if article.status != ArticleStatus.PUBLISHED:
        if not user:
            raise HTTPException(status_code=401, detail="请登录后查看私有文章")
        if not (is_author or is_admin):
            raise HTTPException(status_code=403, detail="无权访问该文章")

    # 使用原子操作增加阅读量，避免并发问题
    await db.execute(
        update(Article)
        .where(Article.id == article_id)
        .values(view_count=Article.view_count + 1)
    )
    await db.commit()
    
    # 刷新对象以获取最新的 view_count
    await db.refresh(article)

    # 查询点赞信息
    # 1. 查询总点赞数
    like_count_res = await db.execute(
        select(func.count(ArticleLike.id)).where(ArticleLike.article_id == article_id)
    )
    article.like_count = like_count_res.scalar() or 0

    # 2. 查询当前用户是否已点赞
    if user:
        like_check = await db.execute(
            select(ArticleLike).where(
                ArticleLike.article_id == article_id,
                ArticleLike.user_id == user.id
            )
        )
        article.is_liked = like_check.scalars().first() is not None
    else:
        article.is_liked = False

    return article


# --- 接口 6：POST /{article_id}/withdraw (撤回审核) ---
@router.post("/{article_id}/withdraw", summary="撤回发布")
async def withdraw_article(article_id: int, user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(Article).where(Article.id == article_id))
    article = res.scalars().first()
    if not article or (article.user_id != user.id and user.role != UserRole.ADMIN):
        raise HTTPException(status_code=403)

    if article.status != ArticleStatus.PENDING:
        raise HTTPException(status_code=400, detail="只有处于待审核状态的文章可以撤回")

    article.status = ArticleStatus.DRAFT
    await db.commit()
    return {"message": "已撤回至草稿状态"}


# --- 接口 7：POST /admin/articles/{article_id}/review (管理员审核) ---
@router.post("/admin/articles/{article_id}/review", summary="【管理员】审核文章")
async def review_article(article_id: int, action: ArticleReviewAction, admin: User = Depends(allow_admin_only),
                         db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(Article).where(Article.id == article_id))
    article = res.scalars().first()
    if not article:
        raise HTTPException(status_code=404, detail="文章不存在")

    if article.status != ArticleStatus.PENDING:
        raise HTTPException(status_code=400, detail="该文章不处于待审核状态")

    if action.pass_audit:
        article.status = ArticleStatus.PUBLISHED
        article.review_remark = None
        article.reviewed_at = datetime.now()
        article.reviewed_by = admin.id
        article.published_at = datetime.now()
    else:
        if not action.remark or not action.remark.strip():
            raise HTTPException(status_code=400, detail="驳回文章必须填写驳回理由")
        article.status = ArticleStatus.DRAFT
        article.reviewed_at = datetime.now()
        article.reviewed_by = admin.id
        article.review_remark = action.remark

    await db.commit()
    return {"message": "审核操作成功"}


# --- 接口 8：GET /admin/pending (待审核列表) ---
@router.get("/admin/pending", summary="【管理员】待审核列表")
async def list_pending_articles(
        page: int = Query(1, ge=1),
        size: int = Query(20, ge=1),
        admin: User = Depends(allow_admin_only),
        db: AsyncSession = Depends(get_db)
):
    stmt = (
        select(Article)
        .where(Article.status == ArticleStatus.PENDING, Article.deleted_at == None)
        .order_by(Article.submitted_at.asc())
    )
    total_res = await db.execute(select(func.count()).select_from(stmt.subquery()))
    total = total_res.scalar() or 0
    res = await db.execute(stmt.offset((page - 1) * size).limit(size))
    return {
        "items": res.scalars().all(),
        "total": total,
        "page": page,
        "size": size,
        "pages": math.ceil(total / size) if total > 0 else 0
    }


# --- 接口 9：DELETE /{article_id} (软删除) ---
@router.delete("/{article_id}", summary="软删除")
async def soft_delete(article_id: int, user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(Article).where(Article.id == article_id))
    article = res.scalars().first()
    if not article:
        raise HTTPException(status_code=404, detail="文章不存在")
    if article.deleted_at:
        raise HTTPException(status_code=400, detail="文章已在回收站中")
    if user.role != UserRole.ADMIN and article.user_id != user.id:
        raise HTTPException(status_code=403, detail="无权删除此文章")

    article.deleted_at = datetime.now()
    await db.commit()
    return {"message": "已入回收站"}


# --- 接口 10：DELETE /{article_id}/hard (硬删除) ---
@router.delete("/{article_id}/hard", summary="硬删除")
async def hard_delete_article(article_id: int, user: User = Depends(get_current_user),
                              db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(Article).where(Article.id == article_id))
    article = res.scalars().first()
    if not article or (article.user_id != user.id and user.role != UserRole.ADMIN):
        raise HTTPException(status_code=403, detail="无权删除此文章")

    await db.delete(article)
    await db.commit()
    return {"message": "文章已永久删除"}


# --- 接口 11：POST /{article_id}/restore (恢复文章) ---
@router.post("/{article_id}/restore", summary="恢复文章")
async def restore_article(article_id: int, user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    stmt = update(Article).where(Article.id == article_id)
    if user.role != UserRole.ADMIN: stmt = stmt.where(Article.user_id == user.id)
    await db.execute(stmt.values(deleted_at=None))
    await db.commit()
    return {"message": "已恢复"}


# --- 接口 12：GET /my/list (我的文章列表) ---
@router.get("/my/list", summary="我的文章列表")
async def get_my_articles(page: int = Query(1, ge=1), size: int = Query(10, ge=1),
                          status: Optional[ArticleStatus] = Query(None), user: User = Depends(get_current_user),
                          db: AsyncSession = Depends(get_db)):
    filters = [Article.user_id == user.id, Article.deleted_at == None]
    if status: filters.append(Article.status == status)
    query = select(Article).where(and_(*filters)).order_by(Article.is_pinned.desc(), Article.created_at.desc())
    total_res = await db.execute(select(func.count()).select_from(query.subquery()))
    total = total_res.scalar() or 0
    res = await db.execute(query.offset((page - 1) * size).limit(size))
    return {"items": res.scalars().all(), "total": total, "page": page, "size": size, "pages": math.ceil(total / size)}


# --- 接口 13：GET /public/list (公开文章列表) ---
@router.get("/public/list", summary="公开文章列表")
async def list_public_articles(page: int = Query(1, ge=1), size: int = Query(10, ge=1),
                               category_id: Optional[int] = None, db: AsyncSession = Depends(get_db)):
    filters = [Article.status == ArticleStatus.PUBLISHED, Article.deleted_at == None]
    if category_id: filters.append(Article.category_id == category_id)
    query = select(Article).where(and_(*filters)).order_by(Article.is_pinned.desc(), Article.created_at.desc())
    total_res = await db.execute(select(func.count()).select_from(query.subquery()))
    total = total_res.scalar() or 0
    res = await db.execute(query.offset((page - 1) * size).limit(size))
    return {"items": res.scalars().all(), "total": total, "page": page, "size": size, "pages": math.ceil(total / size)}


# --- 接口 14：GET /admin/all-articles (全站文章列表) ---
@router.get("/admin/all-articles", summary="【管理员】全站文章列表")
async def list_all_articles_admin(page: int = Query(1, ge=1), size: int = Query(10, ge=1),
                                  show_deleted: bool = Query(False, description="是否显示已删除文章"),
                                  admin: User = Depends(allow_admin_only), db: AsyncSession = Depends(get_db)):
    filters = []
    if not show_deleted:
        filters.append(Article.deleted_at == None)

    query = select(Article).where(and_(*filters)).order_by(Article.is_pinned.desc(), Article.created_at.desc())
    total_res = await db.execute(select(func.count()).select_from(query.subquery()))
    total = total_res.scalar() or 0
    res = await db.execute(query.offset((page - 1) * size).limit(size))
    return {"items": res.scalars().all(), "total": total, "page": page, "size": size, "pages": math.ceil(total / size)}


# --- 接口 15：PUT /admin/articles/{article_id}/pin (管理员置顶/取消置顶) ---
@router.put("/admin/articles/{article_id}/pin", summary="【管理员】置顶/取消置顶文章")
async def toggle_pin_article(
    article_id: int,
    admin: User = Depends(allow_admin_only),
    db: AsyncSession = Depends(get_db)
):
    res = await db.execute(select(Article).where(Article.id == article_id))
    article = res.scalars().first()
    if not article:
        raise HTTPException(status_code=404, detail="文章不存在")
    article.is_pinned = not article.is_pinned  # SQLAlchemy Boolean 类型应使用布尔值
    await db.commit()
    return {"message": "已置顶" if article.is_pinned else "已取消置顶", "is_pinned": article.is_pinned}


# --- 接口 16：GET /public/archive (公开文章归档) ---
@router.get("/public/archive", summary="文章归档（按年月统计）")
async def get_article_archive(db: AsyncSession = Depends(get_db)):
    # 只统计已发布且未删除的文章
    stmt = (
        select(
            func.YEAR(Article.published_at).label("year"),
            func.MONTH(Article.published_at).label("month"),
            func.count(Article.id).label("count")
        )
        .where(Article.status == ArticleStatus.PUBLISHED, Article.deleted_at == None)
        .group_by("year", "month")
        .order_by("year", "month")
    )
    res = await db.execute(stmt)
    rows = res.all()
    return [
        {"year": row.year, "month": row.month, "count": row.count}
        for row in rows
    ]


# --- 接口 17：POST /{article_id}/like (文章点赞/取消点赞) ---
@router.post("/{article_id}/like", summary="文章点赞/取消点赞")
async def toggle_article_like(
        article_id: int,
        user: User = Depends(get_current_user),
        db: AsyncSession = Depends(get_db)
):
    # 校验文章是否存在
    article_res = await db.execute(select(Article).where(Article.id == article_id))
    if not article_res.scalars().first():
        raise HTTPException(status_code=404, detail="文章不存在")

    # 检查是否已点赞
    stmt = select(ArticleLike).where(
        and_(ArticleLike.article_id == article_id, ArticleLike.user_id == user.id)
    )
    res = await db.execute(stmt)
    existing_like = res.scalars().first()

    if existing_like:
        await db.delete(existing_like)
        liked = False
    else:
        new_like = ArticleLike(article_id=article_id, user_id=user.id)
        db.add(new_like)
        liked = True

    try:
        await db.commit()
    except IntegrityError:
        await db.rollback()
        raise HTTPException(status_code=400, detail="操作过于频繁")

    # 获取最新点赞数
    count_res = await db.execute(select(func.count(ArticleLike.id)).where(ArticleLike.article_id == article_id))
    return {"liked": liked, "like_count": count_res.scalar() or 0}


# --- 接口 18：GET /{article_id}/like/count (获取文章点赞数) ---
@router.get("/{article_id}/like/count", summary="获取文章点赞数")
async def get_article_like_count(article_id: int, db: AsyncSession = Depends(get_db)):
    count_res = await db.execute(
        select(func.count(ArticleLike.id)).where(ArticleLike.article_id == article_id)
    )
    return {"article_id": article_id, "like_count": count_res.scalar() or 0}


# --- 辅助函数：从内容中提取摘要 ---
def extract_summary(content: str, max_length: int = 200) -> str:
    """从内容中提取纯文本前max_length字符作为摘要，去掉Markdown语法符号"""
    # 去除常见的Markdown标记
    text = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)  # 标题标记（只删除行首的#）
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 粗体 **text**
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # 斜体 *text*
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # 链接 [text](url)
    text = re.sub(r'!\[(.+?)\]\(.+?\)', r'\1', text)  # 图片 ![alt](url)
    text = re.sub(r'`{1,3}(.+?)`{1,3}', r'\1', text)  # 代码块 `code`
    text = re.sub(r'^\s*>\s*', '', text, flags=re.MULTILINE)  # 引用 >
    text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)  # 列表项
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)  # 有序列表
    
    # 清理多余空白
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = ' '.join(text.split())
    
    return text[:max_length]


# --- 辅助函数：从txt文件提取标题和内容 ---
def parse_txt_file(content: bytes, filename: str) -> tuple:
    """解析txt文件，返回(title, content)"""
    text = content.decode('utf-8', errors='ignore')
    lines = text.split('\n')
    
    # 找第一个长度1-100的非空行作为标题
    title = None
    for line in lines:
        stripped = line.strip()
        if 1 <= len(stripped) <= 100:
            title = stripped
            break
    
    # 如果没有符合条件的行，用文件名（去掉.txt后缀）作为标题
    if not title:
        title = os.path.splitext(filename)[0]
    
    return title, text


# --- 辅助函数：从md文件提取标题和内容 ---
def parse_md_file(content: bytes, filename: str) -> tuple:
    """解析md文件，返回(title, content)"""
    text = content.decode('utf-8', errors='ignore')
    lines = text.split('\n')
    
    # 找第一个以"# "开头的行作为标题
    title = None
    for line in lines:
        if line.startswith('# '):
            title = line[2:].strip()
            break
    
    # 如果没有"# "开头的行，用文件名（去掉.md后缀）作为标题
    if not title:
        title = os.path.splitext(filename)[0]
    
    return title, text


# --- 辅助函数：从docx文件提取标题和内容 ---
def parse_docx_file(file_path: str, filename: str) -> tuple:
    """解析docx文件，返回(title, content)"""
    from docx import Document
    
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = '\n'.join(paragraphs)
    
    # 标题策略：找第一个样式为Heading 1的段落
    title = None
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1' and para.text.strip():
            title = para.text.strip()
            break
    
    # 如果没有Heading 1，找第一个bold=True的段落
    if not title:
        for para in doc.paragraphs:
            for run in para.runs:
                if run.bold and para.text.strip():
                    title = para.text.strip()
                    break
            if title:
                break
    
    # 如果都没有，用文件名（去掉.docx后缀）作为标题
    if not title:
        title = os.path.splitext(filename)[0]
    
    return title, content


# --- 接口 19：POST /admin/import/single (单篇导入文章) ---
@router.post("/admin/import/single", summary="【管理员】单篇导入文章")
async def import_single_article(
    file: UploadFile = File(...),
    admin: User = Depends(allow_admin_only),
    db: AsyncSession = Depends(get_db)
):
    # 文件大小限制10MB
    file_content = await file.read()
    if len(file_content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件大小不能超过10MB")
    
    # 检测文件扩展名
    ext = os.path.splitext(file.filename)[1].lower()
    
    try:
        if ext == '.txt':
            title, content = parse_txt_file(file_content, file.filename)
        elif ext == '.md':
            title, content = parse_md_file(file_content, file.filename)
        elif ext == '.docx':
            # docx需要先保存到临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            try:
                title, content = parse_docx_file(tmp_path, file.filename)
            finally:
                os.unlink(tmp_path)
        else:
            raise HTTPException(status_code=400, detail="不支持的文件格式，仅支持 .txt、.md、.docx")
        
        # 生成摘要
        summary = extract_summary(content, 200)
        
        # 创建Article对象
        article = Article(
            title=title,
            content=content,
            summary=summary,
            user_id=admin.id,
            status=ArticleStatus.DRAFT,
            category_id=None,
            version=1
        )
        db.add(article)
        await db.commit()
        await db.refresh(article)
        
        return {
            "article_id": article.id,
            "title": article.title,
            "message": "文章导入成功，请前往草稿箱编辑"
        }
    except HTTPException:
        raise
    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=f"导入失败: {str(e)}")


# --- 接口 20：POST /admin/import/batch (批量导入文章) ---
@router.post("/admin/import/batch", summary="【管理员】批量导入文章")
async def import_batch_articles(
    files: List[UploadFile] = File(...),
    admin: User = Depends(allow_admin_only),
    db: AsyncSession = Depends(get_db)
):
    # 总大小限制50MB
    # 先读取所有文件大小（不消耗文件流）
    total_size = 0
    for f in files:
        content = await f.read()
        total_size += len(content)
        await f.seek(0)  # 重置指针以便后续处理
    
    if total_size > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="总文件大小不能超过50MB")
    
    success_count = 0
    failed_list = []
    articles_list = []
    
    for file in files:
        try:
            # 读取文件内容
            file_content = await file.read()
            ext = os.path.splitext(file.filename)[1].lower()
            
            # 复用单篇导入的文件解析逻辑
            if ext == '.txt':
                title, content = parse_txt_file(file_content, file.filename)
            elif ext == '.md':
                title, content = parse_md_file(file_content, file.filename)
            elif ext == '.docx':
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    tmp.write(file_content)
                    tmp_path = tmp.name
                try:
                    title, content = parse_docx_file(tmp_path, file.filename)
                finally:
                    os.unlink(tmp_path)
            else:
                failed_list.append({
                    "filename": file.filename,
                    "reason": "不支持的文件格式"
                })
                continue
            
            # 生成摘要
            summary = extract_summary(content, 200)
            
            # 创建Article对象（独立事务）
            article = Article(
                title=title,
                content=content,
                summary=summary,
                user_id=admin.id,
                status=ArticleStatus.DRAFT,
                category_id=None,
                version=1
            )
            db.add(article)
            await db.commit()
            await db.refresh(article)
            
            articles_list.append({
                "article_id": article.id,
                "title": article.title
            })
            success_count += 1
            
        except Exception as e:
            await db.rollback()
            failed_list.append({
                "filename": file.filename,
                "reason": str(e)
            })
            continue
    
    return {
        "total": len(files),
        "success": success_count,
        "failed": failed_list,
        "articles": articles_list
    }


# --- 接口 21：POST /admin/upload-images/batch (批量上传图片) ---
@router.post("/admin/upload-images/batch", summary="【管理员】批量上传图片")
async def batch_upload_images(
    file: UploadFile = File(...),
    admin: User = Depends(allow_admin_only)
):
    # 校验文件扩展名
    ext = os.path.splitext(file.filename)[1].lower()
    if ext != '.zip':
        raise HTTPException(status_code=400, detail="只支持.zip格式的压缩包")
    
    # 文件大小限制50MB
    file_content = await file.read()
    if len(file_content) > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件大小不能超过50MB")
    
    # 支持的图片格式
    allowed_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.webp'}
    
    # 创建临时目录解压
    temp_dir = tempfile.mkdtemp()
    extract_dir = os.path.join(temp_dir, 'extracted')  # 创建子目录用于解压
    os.makedirs(extract_dir, exist_ok=True)
    zip_path = os.path.join(temp_dir, 'upload.zip')
    
    try:
        # 写入zip文件
        with open(zip_path, 'wb') as f:
            f.write(file_content)
        
        # 解压到子目录（避免zip文件本身被遍历）
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
        
        # 遍历解压后的文件
        urls = []
        failed_list = []
        success_count = 0
        
        for root, dirs, files in os.walk(temp_dir):
            for filename in files:
                file_ext = os.path.splitext(filename)[1].lower()
                
                # 只处理图片格式
                if file_ext not in allowed_extensions:
                    failed_list.append({
                        "filename": filename,
                        "reason": "不是图片格式"
                    })
                    continue
                
                try:
                    # 生成唯一文件名
                    unique_name = f"{uuid.uuid4().hex}{file_ext}"
                    dest_path = os.path.join(IMAGE_STORAGE, unique_name)
                    src_path = os.path.join(root, filename)
                    
                    # 确保目标目录存在
                    os.makedirs(IMAGE_STORAGE, exist_ok=True)
                    
                    # 移动到storage/images目录
                    shutil.move(src_path, dest_path)
                    
                    urls.append(f"/storage/images/{unique_name}")
                    success_count += 1
                    
                except Exception as e:
                    failed_list.append({
                        "filename": filename,
                        "reason": str(e)
                    })
                    continue
        
        return {
            "total": success_count + len(failed_list),
            "success": success_count,
            "failed": failed_list,
            "urls": urls
        }
        
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="无效的zip文件")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")
    finally:
        # 清理临时目录
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except:
            pass